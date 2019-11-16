import calendar
import datetime
import os
import uuid

from django.conf import settings
from django.db.models import Q, Sum

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .models import MonthlyFieldService
from publishers.models import Group


def compute_month_year(date):
    default_month_year = '{}-{}'.format(date.year, date.month)
    return default_month_year


def get_months_and_years(date_from, date_to):
    """
    Return months and years from [date from] and [date to]
    """
    date_from = date_from.split('-')
    date_to = date_to.split('-')
    df = '{}-{}'.format(date_from[0], date_from[1])
    dt = '{}-{}'.format(date_to[0], date_to[1])

    return {
        'fm': date_from[1],
        'tm': date_to[1],
        'fy': date_from[0],
        'ty': date_to[0],
        'df': df,
        'dt': dt,
    }


def aggregate_mfs_queryset(queryset):
    q = queryset.aggregate(
        Sum('placements'),
        Sum('video_showing'),
        Sum('hours'),
        Sum('return_visits'),
        Sum('bible_study'),
        rp_placements=Sum('placements', filter=Q(
            pioneering__pioneer_type='RP')),
        rp_video_showing=Sum('video_showing', filter=Q(
            pioneering__pioneer_type='RP')),
        rp_hours=Sum('hours', filter=Q(pioneering__pioneer_type='RP')),
        rp_return_visits=Sum('return_visits', filter=Q(
            pioneering__pioneer_type='RP')),
        rp_bible_study=Sum('bible_study', filter=Q(
            pioneering__pioneer_type='RP')),
        au_placements=Sum('placements', filter=Q(
            pioneering__pioneer_type='AU')),
        au_video_showing=Sum('video_showing', filter=Q(
            pioneering__pioneer_type='AU')),
        au_hours=Sum('hours', filter=Q(pioneering__pioneer_type='AU')),
        au_return_visits=Sum('return_visits', filter=Q(
            pioneering__pioneer_type='AU')),
        au_bible_study=Sum('bible_study', filter=Q(
            pioneering__pioneer_type='AU')),
        sp_placements=Sum('placements', filter=Q(
            pioneering__pioneer_type='SP')),
        sp_video_showing=Sum('video_showing', filter=Q(
            pioneering__pioneer_type='SP')),
        sp_hours=Sum('hours', filter=Q(pioneering__pioneer_type='SP')),
        sp_return_visits=Sum('return_visits', filter=Q(
            pioneering__pioneer_type='SP')),
        sp_bible_study=Sum('bible_study', filter=Q(
            pioneering__pioneer_type='SP')),
    )

    return q


def get_mfs_data(date_from, date_to, pk, is_publisher=True):
    months_years = get_months_and_years(date_from, date_to)
    from_month, from_year = months_years['fm'], months_years['fy']
    to_month, to_year = months_years['tm'], months_years['ty']

    queryset = MonthlyFieldService.objects.filter(
        month_ending__month__gte=from_month,
        month_ending__year__gte=from_year
    ).order_by(
        '-month_ending', 'publisher__last_name', 'publisher__first_name')

    queryset = queryset.select_related('publisher', 'pioneering')

    if is_publisher:
        queryset = queryset.filter(publisher=pk)

    queryset = queryset.filter(
        month_ending__month__lte=to_month,
        month_ending__year__lte=to_year
    )

    if not is_publisher:
        # filter for group only
        group = Group.objects.get(pk=pk)
        queryset = queryset.filter(group=pk)

    totals = aggregate_mfs_queryset(queryset)

    data = {
        'queryset': queryset,
        'totals': totals,
        'from': '{}-{}'.format(from_month, from_year),
        'to': '{}-{}'.format(to_month, to_year),
    }

    if not is_publisher:
        data['group'] = group

    return data


def generate_mfs(data, report_type='group'):
    title = 'Monthly Field Service Report'
    doc = Document()

    doc.add_heading(
        title, 0).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if report_type == 'group':
        p = doc.add_paragraph(data['group'])
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = 'Congregation: {}'.format(data['congregation'])
    p = doc.add_paragraph(p)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph('For the month {}'.format(data['month']))
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    reports = data['queryset']

    headers = [
        ('Month', 'Month'),
        ('Name', 'Name'),
        ('P', 'Placements'),
        ('VS', 'Video Showings'),
        ('H', 'Hours'),
        ('RV', 'Return Visits'),
        ('BS', 'Bible Studies'),
        ('Comments', 'Comments')
    ]

    table = doc.add_table(rows=1, cols=len(headers))

    hdr = table.rows[0].cells
    for i in range(len(headers)):
        label = headers[i][0]
        hdr[i].text = label

    for report in reports:
        row_cells = table.add_row().cells
        row_cells[0].text = report.month_ending.strftime('%b-%Y')
        row_cells[1].text = str(report.publisher.name)
        row_cells[2].text = str(report.placements)
        row_cells[3].text = str(report.video_showing)
        row_cells[4].text = str(report.hours)
        row_cells[5].text = str(report.return_visits)
        row_cells[6].text = str(report.bible_study)
        row_cells[7].text = str(report.comments)

    # totals
    totals = data['totals']
    row_cells = table.add_row().cells
    row_cells[0].text = 'TOTALS'
    row_cells[2].text = str(totals['placements__sum'])
    row_cells[3].text = str(totals['video_showing__sum'])
    row_cells[4].text = str(totals['hours__sum'])
    row_cells[5].text = str(totals['return_visits__sum'])
    row_cells[6].text = str(totals['bible_study__sum'])

    # convert none values to 0
    for (k, v) in totals.items():
        if v is None:
            totals[k] = 0

    if report_type == 'group':
        other_rows = [
            ('rp', 'Regular Pioneer'),
            ('au', 'Auxillary Pioneer'),
            ('sp', 'Special Pioneer')]
        for r in other_rows:
            row = r[0]
            placements = totals['{}_placements'.format(row)]
            video_showing = totals['{}_video_showing'.format(row)]
            hours = totals['{}_hours'.format(row)]
            return_visits = totals['{}_return_visits'.format(row)]
            bible_study = totals['{}_bible_study'.format(row)]

            row_cells = table.add_row().cells
            row_cells[0].text = r[1].upper()
            row_cells[2].text = str(placements)
            row_cells[3].text = str(video_showing)
            row_cells[4].text = str(hours)
            row_cells[5].text = str(return_visits)
            row_cells[6].text = str(bible_study)

            totals['placements__sum'] -= placements
            totals['video_showing__sum'] -= video_showing
            totals['hours__sum'] -= hours
            totals['return_visits__sum'] -= return_visits
            totals['bible_study__sum'] -= bible_study

    # publishers row
    row_cells = table.add_row().cells
    row_cells[0].text = 'PUBLISHERS'
    row_cells[2].text = str(totals['placements__sum'])
    row_cells[3].text = str(totals['video_showing__sum'])
    row_cells[4].text = str(totals['hours__sum'])
    row_cells[5].text = str(totals['return_visits__sum'])
    row_cells[6].text = str(totals['bible_study__sum'])

    filename = '{}.docx'.format(str(uuid.uuid1()))
    fullpath = os.path.join(settings.ROOT_DIR, 'media')
    fullpath = os.path.join(fullpath, 'temp--mfs_export')
    if not os.path.exists(fullpath):
        os.mkdir(fullpath)
    fullpath = os.path.join(fullpath, filename)
    doc.save(fullpath)

    return {
        'doc': doc,
        'filename': filename,
        'fullpath': fullpath,
    }


def mfs_stats(month, year):
    q = MonthlyFieldService.objects.filter(month_ending__month=month)
    q = q.filter(month_ending__year=year)
    q = q.aggregate(Sum('hours'), Sum('return_visits'), Sum('bible_study'))

    hours = q['hours__sum']
    return_visits = q['return_visits__sum']
    bible_studies = q['bible_study__sum']
    stats = {
        'hours': hours,
        'return_visits': return_visits,
        'bible_studies': bible_studies
    }
    return stats


def get_previous_month_end(now=None):
    if now is None:
        now = datetime.datetime.now().date()

    # get previous month
    previous_month = now.month - 1
    previous_year = now.year
    if previous_month == 0:
        previous_month = 12
        previous_year -= 1

    cal = calendar.Calendar()
    month = cal.monthdayscalendar(previous_year, previous_month)

    # previous month end
    month = datetime.date(previous_year, previous_month, max(month[-1]))
    return month
